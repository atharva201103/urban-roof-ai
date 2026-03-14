import os
from docx import Document

def extract_text_from_docx(docx_path):
    """
    Extracts text from a given .docx file.
    
    Args:
        docx_path (str): The path to the input .docx file.
        
    Returns:
        str: The extracted text as a single string.
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # You can also extract text from tables if needed
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
                        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""
